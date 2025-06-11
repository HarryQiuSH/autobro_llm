"""
Document Processing Module for RAG Bot

This module handles document upload, processing, and text extraction
for various file formats including PDF, DOCX, and TXT files.

Author: Harry Qiu
"""

import os
import tempfile
from typing import List, Dict, Any
import streamlit as st
from io import BytesIO

# Document processing imports
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

# Google Drive imports
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2 import service_account


class DocumentProcessor:
    """Handles document processing and text extraction."""

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def extract_text_from_pdf(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            return text
        except Exception as e:
            st.error(f"Error processing PDF {filename}: {str(e)}")
            return ""

    def extract_text_from_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing DOCX {filename}: {str(e)}")
            return ""

    def extract_text_from_txt(self, file_content: bytes, filename: str) -> str:
        """Extract text from TXT file."""
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return file_content.decode("latin-1")
            except Exception as e:
                st.error(f"Error processing TXT {filename}: {str(e)}")
                return ""

    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Process an uploaded file and extract text."""
        file_content = uploaded_file.read()
        filename = uploaded_file.name
        file_type = filename.split(".")[-1].lower()

        # Extract text based on file type
        if file_type == "pdf":
            text = self.extract_text_from_pdf(file_content, filename)
        elif file_type == "docx":
            text = self.extract_text_from_docx(file_content, filename)
        elif file_type == "txt":
            text = self.extract_text_from_txt(file_content, filename)
        else:
            st.error(f"Unsupported file type: {file_type}")
            return None

        if not text.strip():
            st.error(f"No text could be extracted from {filename}")
            return None

        return {
            "filename": filename,
            "content": text,
            "file_type": file_type,
            "size": len(file_content),
        }

    def chunk_documents(
        self, documents: List[Dict[str, Any]]
    ) -> List[LangchainDocument]:
        """Split documents into chunks for vector storage."""
        chunked_docs = []

        for doc in documents:
            # Split the text into chunks
            chunks = self.text_splitter.split_text(doc["content"])

            # Create LangChain Document objects with metadata
            for i, chunk in enumerate(chunks):
                chunked_doc = LangchainDocument(
                    page_content=chunk,
                    metadata={
                        "filename": doc["filename"],
                        "file_type": doc["file_type"],
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                    },
                )
                chunked_docs.append(chunked_doc)

        return chunked_docs

    def display_document_info(self, documents: List[Dict[str, Any]]):
        """Display information about processed documents."""
        if not documents:
            st.info("No documents uploaded yet.")
            return

        st.subheader("📄 Uploaded Documents")
        for doc in documents:
            with st.expander(f"{doc['filename']} ({doc['file_type'].upper()})"):
                st.write(f"**Size:** {doc['size']:,} bytes")
                st.write(f"**Content Preview:**")
                preview = (
                    doc["content"][:500] + "..."
                    if len(doc["content"]) > 500
                    else doc["content"]
                )
                st.text(preview)


class GoogleDriveProcessor:
    """Handles Google Drive integration for document retrieval."""

    def __init__(self, api_key: str, folder_id: str):
        self.api_key = api_key
        self.folder_id = folder_id
        self.service = None
        self._initialize_service()

    def _initialize_service(self):
        """Initialize Google Drive service."""
        try:
            self.service = build("drive", "v3", developerKey=self.api_key)
        except Exception as e:
            st.error(f"Failed to initialize Google Drive service: {str(e)}")

    def list_files_in_folder(self) -> List[Dict[str, str]]:
        """List files in the specified Google Drive folder."""
        if not self.service:
            return []

        try:
            query = f"'{self.folder_id}' in parents and trashed=false"
            results = (
                self.service.files()
                .list(q=query, fields="files(id, name, mimeType, size)")
                .execute()
            )

            files = results.get("files", [])

            # Filter for supported file types
            supported_types = [
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "text/plain",
            ]

            filtered_files = [
                file for file in files if file.get("mimeType") in supported_types
            ]

            return filtered_files

        except Exception as e:
            st.error(f"Error listing Google Drive files: {str(e)}")
            return []

    def download_file(self, file_id: str, filename: str) -> bytes:
        """Download a file from Google Drive."""
        if not self.service:
            return None

        try:
            request = self.service.files().get_media(fileId=file_id)
            file_content = BytesIO()
            downloader = MediaIoBaseDownload(file_content, request)

            done = False
            while done is False:
                status, done = downloader.next_chunk()

            file_content.seek(0)
            return file_content.read()

        except Exception as e:
            st.error(f"Error downloading file {filename}: {str(e)}")
            return None
